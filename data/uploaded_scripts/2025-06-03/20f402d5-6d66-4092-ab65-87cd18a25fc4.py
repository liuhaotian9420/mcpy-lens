"""对于文档类文件实现关键词匹配的逻辑"""

import os
import json
import re
import zipfile
import tempfile
from abc import ABC, abstractmethod
from datetime import date, datetime
from pypdf import PdfReader
from docx import Document


class CorruptedDocument(Exception):
    def __init__(self, message,file_path: str):
        self.message = message
        self.file_path = file_path
    def __str__(self):
        return f"Corrupted document: {self.message}, file: {self.file_path}"
    
class RuleError(Exception):
    def __init__(self, message, key):
        self.message = message
        self.key = key
    def __str__(self):
        return f"Rule error: {self.message}, key: {self.key}"
    



class ComplexEncoder(json.JSONEncoder):
    """
    JSONEncoder 重写，使 datetime 和 date 类型可以被 JSON 序列化
    """

    def default(self, o):
        if isinstance(o, datetime):
            return o.strftime("%Y-%m-%d %H:%M:%S")
        elif isinstance(o, date):
            return o.strftime("%Y-%m-%d")
        else:
            return json.JSONEncoder.default(self, o)

class BaseDocumentReader(ABC):
    """
    文档读取
    """

    def __init__(
        self,
    ):
        self.meta_data = {}

    @abstractmethod
    def read_document(self, file_path: str) -> str:
        """
        Reads the document content from the given file path
        Returns:
        str: The document content
        """

    @abstractmethod
    def get_metadata(self) -> dict:
        """
        Returns:
        dict: Metadata extracted from the document
        """
        return json.dumps(
            self.meta_data, indent=4, cls=ComplexEncoder, ensure_ascii=False
        )

class PDFReader(BaseDocumentReader):
    """
    读取 .pdf 文件，不支持扫描件的读取
    """

    def __init__(
        self,
    ):
        super().__init__()
        self.pdf_path = None
        self.text = ""

    def read_document(self, file_path: str) -> str:
        reader = PdfReader(file_path)
        # fetch metadata
        self.meta_data = {
            "author": reader.metadata.author,
            "title": reader.metadata.title,
            "subject": reader.metadata.subject,
            "creator": reader.metadata.creator,
            "producer": reader.metadata.producer,
            "creation_date": reader.metadata.creation_date,
            "modification_date": reader.metadata.modification_date,
            "pages": reader.get_num_pages(),
        }
        # fetch text
        self.text = " ".join(page.extract_text() or "" for page in reader.pages)
        return self.text

    def get_metadata(self):
        return super().get_metadata()

class DocxReader(BaseDocumentReader):
    """
    读取 docx 文件（Words 2007+)
    """

    def __init__(
        self,
    ):
        super().__init__()
        self.docx_path = None
        self.text = ""

    def read_document(self, file_path):
        doc = Document(file_path)
        core_properties = doc.core_properties
        self.meta_data = {
            "author": core_properties.author,
            "title": core_properties.title,
            "subject": core_properties.subject,
            "creator": core_properties.author,
            "producer": core_properties.last_modified_by,
            "creation_date": core_properties.created,
            "modification_date": core_properties.modified,
            "pages": len(doc.paragraphs),
        }
        return "\n".join(para.text for para in doc.paragraphs)

    def get_metadata(self):
        return super().get_metadata()

class IpynbReader(BaseDocumentReader):
    """
    读取 ipynb 文件
    """

    def __init__(
        self,
    ):
        super().__init__()
        self.ipynb_path = None
        self.text = ""
        self.code = ""
        self.meta_data = {}

    def read_document(self, file_path):
        '''
        读取 ipynb 文件并提取文本和代码
        '''
        ipynb_json = json.load(open(file_path, "r", encoding="utf-8"))
        metadata = ipynb_json.get("metadata", {})
        self.meta_data = metadata
        cells = ipynb_json.get("cells", [])
        self.text = "\n".join(
            [' '.join(cell["source"]) for cell in cells if cell["cell_type"] == "markdown"]
        )
        self.code = "\n".join(
            [''.join(cell["source"]) for cell in cells if cell["cell_type"] == "code"]
        )
        return self.text
    
    def get_metadata(self):
        return super().get_metadata()

# ------- Helper Functions ---
def is_valid_file(file_path, file_extensions: str) -> bool:
    """
    检查文件的类型是否符合要求
    """
    if isinstance(file_extensions, str):
        file_extensions = [file_extensions]
    # checks if file_path ends with one of the allowed extensions
    passes_checks = any(file_path.lower().endswith(ext) for ext in file_extensions)
    return os.path.isfile(file_path) and passes_checks


def validate_and_extract_zip(zip_path):
    """
    Extract files from a zip file and validate their extensions
    Returns:
    list: Valid file paths in the zip file
    """
    with zipfile.ZipFile(zip_path, "r") as zip_ref:
        file_list = zip_ref.namelist()
        allowed_ext = (".pdf", ".docx",".document",".ipynb")
        for file in file_list:
            if not file.lower().endswith(allowed_ext):
                raise ValueError(f"Unsupported file in zip: {file}")
        temp_dir = tempfile.mkdtemp()
        zip_ref.extractall(temp_dir)
        return [os.path.join(temp_dir, f) for f in file_list]


def match_rules(path, rules: dict) -> dict:
    """
    匹配 path 最接近的 rules
    """
    normalized_rules = {k.lower(): rules[k] for k in rules}
    escaped_keys = [re.escape(k.lower()) for k in rules]
    matches = re.findall(r"|".join(escaped_keys), path.lower())
    if matches:
        best_match = max(matches, key=len)  # choose longest match
        return normalized_rules.get(best_match, {})
    else:
        return {}


# --- Scoring ---
def score_document(text, rule):
    """
    基于 rule 的分数计算
    """
    keywords = rule.get("keywords", [])
    scores = rule.get("scores", [])
    max_score = rule.get("max_score", 0)
    match_mode = rule.get("match_mode", "regex").lower()

    if len(keywords) != len(scores):
        raise ValueError("Keywords and scores must be of same length")

    final_score = max_score
    missing = []

    for keyword, keyword_score in zip(keywords, scores):
        try:
            found = (
                keyword.lower() in text.lower()
                if match_mode == "exact"
                else re.search(keyword, text, re.IGNORECASE)
            )
        except re.error:
            raise Exception(f"Invalid regex: {keyword}")
            # found = False  # Invalid regex
        if not found:
            final_score -= keyword_score
            missing.append(keyword)

    return max(final_score, 0), missing

# --- Helper for single document scoring ---
def process_single_file(file_path, rules, doc_type):
    """
    单个文件的分数计算
    """

    rule = match_rules(file_path, rules)
    if not rule:
        return {"message": "No matching rule for file", "score": 0, "extension": []}

    try:
        if doc_type == ".pdf":
            reader = PDFReader()
            text = reader.read_document(file_path)
        elif doc_type == ".docx" or doc_type==".document":
            reader = DocxReader()
            text = reader.read_document(file_path)
        elif doc_type == ".ipynb":
            reader = IpynbReader()
            text = reader.read_document(file_path)
        else:
            return {
                "message": f"Unsupported file: {file_path}",
                "score": 0,
                "extension": [],
            }
    except Exception as e:
        raise Exception(f"Reading file failed: {e}")
    try:
        score_val, missing = score_document(text, rule)
        return {"message": "success", "score": score_val, "extension": missing}
    except Exception as e:
        return {"message": f"Scoring failed: {e}", "score": 0, "extension": []}
    
# --- Main Handler ---
def score(submit_file=None, test_file=None,options=None):
    """
    基于 json_path 的规则匹配，并且对于 doc_path 进行分数计算
    """
    doc_path = submit_file
    json_path = test_file

    if not is_valid_file(doc_path, [".pdf", ".docx", ".zip",".document",".ipynb"]):
        return {"message": "document path is not founc", "score": 0, "extension": [],"code":-100}
    if not is_valid_file(json_path, [".json"]):
        return {"message": "config json path is not found", "score": 0, "extension": [],"code":-101}

    doc_type = os.path.splitext(doc_path)[1].lower()

    if doc_type == ".zip":
        try:
            paths = validate_and_extract_zip(doc_path)
        except Exception as e:
            return {
                "message": f"Zip validation failed: {e}",
                "score": 0,
                "extension": [],
                "code":-102
            }
    else:
        paths = [doc_path]

    json_rules = json.load(open(json_path, "r", encoding="utf-8"))

    scores = []
    all_missing = []
    for path in paths:
        result = process_single_file(path, json_rules, doc_type)
        if result["message"] == "success":
            scores.append(result["score"])
            all_missing.extend(result["extension"])

    if not scores:
        return {"message": "No valid documents processed", "score": 0, "extension": [],"code":0}

    avg_score = round(sum(scores) / len(scores), 2)
    return {
        "message": "success",
        "score": avg_score,
        "code":0,
        "extension": list(set(all_missing)),
    }
